import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = docx.Document()

# Page margins: 0.35 in
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)

# Set base font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)
font.color.rgb = RGBColor(0x22, 0x22, 0x22)

def add_header(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    
    run = p.add_run("BENOUALI ABDELKADER YAHIA ZAKARIA\n")
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    run_sub = p.add_run("Master's Student in AI & Cybersecurity  |  AI Software Engineering Intern Candidate\n")
    run_sub.font.size = Pt(10.5)
    run_sub.bold = True
    
    run_contact = p.add_run("Ain Temouchent, Algeria  |  +213 781 306 713  |  abdelkaderbenouali301@gmail.com\n")
    run_contact.font.size = Pt(9.5)
    
    run_links = p.add_run("LinkedIn: linkedin.com/in/benouali-abdelkader-yahia-zakaria-4a917a247  |  GitHub: github.com/benoualiabdelkader\n")
    run_links.font.size = Pt(9.5)
    
    run_note = p.add_run("Open to Relocation & Remote Roles")
    run_note.italic = True
    run_note.font.size = Pt(9.5)

def add_section_title(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    # Add a thin bottom border / line equivalent
    p_border = doc.add_paragraph()
    p_border.paragraph_format.space_before = Pt(0)
    p_border.paragraph_format.space_after = Pt(3)
    r_line = p_border.add_run("―" * 65)
    r_line.font.size = Pt(6)
    r_line.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

def add_bullet(doc, bold_prefix, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.line_spacing = 1.05
    
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.bold = True
    r_text = p.add_run(text)

add_header(doc)

# SUMMARY
add_section_title(doc, "Professional Summary")
p_sum = doc.add_paragraph("Master's student in AI & Cybersecurity (Ranked 3rd in CS cohort) proficient in Python, C/C++, PyTorch, and Scikit-learn. Demonstrated track record in building end-to-end AI pipelines, model quantization, and deploying cloud microservices with strict privacy and security controls. Seeking the AI Software Engineering Graduate Intern role at Intel to optimize AI frameworks and drive hardware-software co-optimization.")
p_sum.paragraph_format.space_after = Pt(4)
p_sum.paragraph_format.line_spacing = 1.05

# TECHNICAL SKILLS
add_section_title(doc, "Technical Skills")
add_bullet(doc, "AI & Machine Learning: ", "PyTorch, Scikit-learn, Model Quantization, Model Optimization & Fine-Tuning, LLM Pipelines, Sensor & Time-Series Data Modeling")
add_bullet(doc, "Languages & Core CS: ", "Python, C, C++, Data Structures & Algorithms, REST APIs, SQL, PostgreSQL")
add_bullet(doc, "Systems & Hardware Integration: ", "Hardware-Software Integration, Docker Containerization, Linux/Windows Administration, Streamlit, Git/GitHub")
add_bullet(doc, "Research & Benchmarking: ", "Applied AI Research, Model Benchmarking, LaTeX Documentation, Scientific Writing, Software Debugging")

# EDUCATION
add_section_title(doc, "Education")
p_edu = doc.add_paragraph()
p_edu.paragraph_format.space_after = Pt(2)
r_sch = p_edu.add_run("University Ain Temouchent Belhadj Bouchaib, Algeria\n")
r_sch.bold = True
r_m = p_edu.add_run("Master of Science in Computer & Information Sciences — AI & Cybersecurity  (Sept 2025 – July 2027 Expected)\n")
r_m.bold = True
r_b = p_edu.add_run("Bachelor of Science in Computer Science  (Sept 2022 – June 2025 Completed)")
add_bullet(doc, "Academic Distinction: ", "Ranked 3rd out of cohort across L1 & L2 semesters. Core Coursework: C Programming, Data Structures, Algorithms, Computer Architecture, Operating Systems.")

# PROJECTS
add_section_title(doc, "AI & Software Engineering Projects")

# WriteLens
p_p1 = doc.add_paragraph()
p_p1.paragraph_format.space_before = Pt(3)
p_p1.paragraph_format.space_after = Pt(1)
r_p1_title = p_p1.add_run("Adaptive Blended Assessment (WriteLens V2)  |  2025 – 2026\n")
r_p1_title.bold = True
r_p1_url = p_p1.add_run("github.com/benoualiabdelkader/Adaptive-Blended-Assessment-  |  ")
r_p1_url.font.size = Pt(9)
r_p1_stack = p_p1.add_run("Tech Stack: Python, Node.js, Express, React, TypeScript, Docker")
r_p1_stack.italic = True
r_p1_stack.font.size = Pt(9)

add_bullet(doc, "", "Engineered an AI-driven academic writing platform processing 480+ student submissions to generate diagnostic profiles, risk bands, and cohort summaries for 60 students.")
add_bullet(doc, "", "Built a privacy-first Python AI engine integrating Groq LLM API with strict tokenization, fail-closed security, and provenance tracking across Render cloud deployments.")
add_bullet(doc, "", "Developed a REST API backend (Node.js/Express) serving JSON run bundles to a React TypeScript dashboard, containerized via Docker with release automation.")

# EdgeQuant
p_p2 = doc.add_paragraph()
p_p2.paragraph_format.space_before = Pt(4)
p_p2.paragraph_format.space_after = Pt(1)
r_p2_title = p_p2.add_run("EdgeQuant — High-Speed LLM Quantization Toolkit  |  2026\n")
r_p2_title.bold = True
r_p2_url = p_p2.add_run("github.com/benoualiabdelkader/EdgeQuant  |  ")
r_p2_url.font.size = Pt(9)
r_p2_stack = p_p2.add_run("Tech Stack: Python, C++, PyTorch, CUDA, ONNX")
r_p2_stack.italic = True
r_p2_stack.font.size = Pt(9)

add_bullet(doc, "", "Developed an open-source quantization toolkit compressing 7B and 13B parameter Transformer models into 4-bit and 8-bit precision for edge devices and GPUs.")
add_bullet(doc, "", "Optimized inference performance by reducing GPU VRAM footprint by 62% while preserving 98.4% of baseline FP16 model accuracy.")

# EcoSentinel
p_p3 = doc.add_paragraph()
p_p3.paragraph_format.space_before = Pt(4)
p_p3.paragraph_format.space_after = Pt(1)
r_p3_title = p_p3.add_run("EcoSentinel AI — Aquatic Sensor Quality & ML Pipeline  |  2026\n")
r_p3_title.bold = True
r_p3_url = p_p3.add_run("github.com/benoualiabdelkader/eco-sentinel-project  |  ")
r_p3_url.font.size = Pt(9)
r_p3_stack = p_p3.add_run("Tech Stack: Python, Scikit-learn, Pandas, NumPy, Streamlit, LaTeX")
r_p3_stack.italic = True
r_p3_stack.font.size = Pt(9)

add_bullet(doc, "", "Built an environmental monitoring system processing aquatic sensor data (turbidity, dissolved oxygen) to detect water pollution in real time.")
add_bullet(doc, "", "Trained a Scikit-learn SVM classifier (RBF kernel) and deployed an interactive Streamlit simulation web app featuring real-time parameter sliders and published a formal LaTeX research report.")

# LEADERSHIP
add_section_title(doc, "Leadership & Global Programs")
add_bullet(doc, "Aspire Leaders Program Participant — Aspire Institute (Harvard-affiliated) (July 2026 – Present): ", "Selected for a global leadership program; completing coursework and masterclasses on strategic decision-making and technical leadership.")
add_bullet(doc, "Participant — Erasmus+ Virtual Exchange — VIRTUALLYEDU (EU-Funded) (July 2026 – Present): ", "Engaging in an EU-funded virtual exchange focused on digital skills, cybersecurity defense strategies, and cross-cultural technical collaboration.")

# CERTIFICATIONS
add_section_title(doc, "Certifications")
add_bullet(doc, "Google AI Essentials Certification: ", "Generative AI, Responsible AI, Data Ethics")
add_bullet(doc, "Google IT Support Professional Certificate: ", "Linux System Administration, Computer Networking, Hardware Troubleshooting")
add_bullet(doc, "Algerian-American Summer University 2025: ", "Advanced Computing Track (AI, Machine Learning, Cybersecurity & Ethics)")

# LANGUAGES
add_section_title(doc, "Languages")
p_lang = doc.add_paragraph("Arabic: Native  |  English: B2 Level (Intermediate / Working Proficiency)  |  French: B1 Level (Working Proficiency)")
p_lang.paragraph_format.space_before = Pt(2)

doc.save("Benouali_Abdelkader_yahia_zakaria_CV.docx")
print("Successfully generated Benouali_Abdelkader_yahia_zakaria_CV.docx!")
